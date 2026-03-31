import os
import tempfile

import pytest
from docx import Document as DocxDocument

from ingestion.docx_loader import DocxLoader




@pytest.fixture
def docx_loader():
    return DocxLoader()


def create_docx(paragraphs=None, tables=None):
    doc = DocxDocument()

    # Add paragraphs
    if paragraphs:
        for p in paragraphs:
            doc.add_paragraph(p)

    # Add tables
    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = len(table_data[0])
            table = doc.add_table(rows=rows, cols=cols)

            for i, row in enumerate(table_data):
                for j, val in enumerate(row):
                    table.cell(i, j).text = val

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_file.name)
    tmp_file.close()

    return tmp_file.name


def test_load_paragraphs_only(docx_loader):
    file_path = create_docx(paragraphs=["Hello", "World"])

    result = docx_loader.load(file_path)

    assert "Hello" in result
    assert "World" in result

    os.remove(file_path)


def test_load_tables_only(docx_loader):
    file_path = create_docx(
        tables=[
            [["A", "B"], ["1", "2"]]
        ]
    )

    result = docx_loader.load(file_path)

    assert "A | B" in result
    assert "1 | 2" in result

    os.remove(file_path)


def test_load_paragraphs_and_tables(docx_loader):
    file_path = create_docx(
        paragraphs=["Intro"],
        tables=[[["A", "B"], ["1", "2"]]]
    )

    result = docx_loader.load(file_path)

    assert "Intro" in result
    assert "A | B" in result
    assert "1 | 2" in result

    os.remove(file_path)


def test_empty_document(docx_loader):
    file_path = create_docx()

    result = docx_loader.load(file_path)

    assert result == ""

    os.remove(file_path)
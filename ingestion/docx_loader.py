from pathlib import Path

from docx import Document as DocxDocument

from ingestion.base_doc_loader import BaseLoader


class DocxLoader(BaseLoader):
    def __init__(self):
        super().__init__(".docx")

    @staticmethod
    def _extract_paragraphs(doc: DocxDocument) -> list[str]:
        return [
            p.text.strip()
            for p in doc.paragraphs
            if p.text and p.text.strip()
        ]

    @staticmethod
    def _extract_tables(doc: DocxDocument) -> list[str]:
        tables_text = []

        for table in doc.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text and cell.text.strip()
                ]

                if cells:
                    tables_text.append(" | ".join(cells))

        return tables_text

    def load(self, file_path: str) -> str:
        fpath = Path(file_path)
        if not fpath.exists() or not fpath.is_file() or not fpath.suffix == ".docx":
            raise FileNotFoundError(f"{fpath} is not a docx file")
        doc = DocxDocument(file_path)

        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)

        content_parts = []

        if paragraphs:
            content_parts.append("\n".join(paragraphs))

        if tables:
            content_parts.append("\n".join(tables))

        return "\n\n".join(content_parts)